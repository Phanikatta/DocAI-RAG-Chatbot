import io
import logging
from pathlib import Path
from typing import List, Dict

logger = logging.getLogger(__name__)

# ── OCR availability check ────────────────────────────────────────────────────
OCR_AVAILABLE = False
try:
    import pytesseract
    pytesseract.get_tesseract_version()
    OCR_AVAILABLE = True
    logger.info("Tesseract OCR available.")
except Exception:
    logger.warning(
        "Tesseract not found — scanned PDF pages will be skipped. "
        "Install Tesseract: https://github.com/UB-Mannheim/tesseract/wiki"
    )


def _ocr_image(pil_image) -> str:
    try:
        import pytesseract
        return pytesseract.image_to_string(pil_image, lang="eng").strip()
    except Exception as e:
        logger.warning(f"OCR failed: {e}")
        return ""


# ── PDF parsing (PyMuPDF — no Poppler needed) ─────────────────────────────────
def extract_text_from_pdf(file_path: Path) -> List[Dict]:
    import fitz  # PyMuPDF

    pages = []
    doc = fitz.open(str(file_path))

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")

        if text and len(text.strip()) > 50:
            pages.append({"page": page_num + 1, "text": text.strip(), "method": "pymupdf"})
        elif OCR_AVAILABLE:
            # Render page to image at 2× for better OCR accuracy
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            from PIL import Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            ocr_text = _ocr_image(img)
            if ocr_text:
                pages.append({"page": page_num + 1, "text": ocr_text, "method": "ocr"})
            else:
                logger.debug(f"Page {page_num + 1}: no text extracted (skipping)")
        else:
            logger.debug(f"Page {page_num + 1}: no text & OCR unavailable (skipping)")

    doc.close()
    return pages


# ── DOCX parsing ──────────────────────────────────────────────────────────────
def extract_text_from_docx(file_path: Path) -> List[Dict]:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(file_path))
    sections: List[Dict] = []
    buffer: List[str] = []
    page = 1

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            buffer.append(text)

        # Flush every ~30 paragraphs as a logical "page"
        if len(buffer) >= 30 or (i == len(doc.paragraphs) - 1 and buffer):
            sections.append({"page": page, "text": "\n".join(buffer), "method": "docx"})
            buffer = []
            page += 1

    # Extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                rows.append(row_text)
        if rows:
            sections.append({"page": page, "text": "\n".join(rows), "method": "docx_table"})
            page += 1

    return sections


# ── Public entry point ────────────────────────────────────────────────────────
def parse_document(file_path: Path) -> List[Dict]:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")
